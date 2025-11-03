from docx import Document
import os
import subprocess
import platform
try:
    from docx2pdf import convert as docx2pdf_convert
except ImportError:
    docx2pdf_convert = None

# Adicionar suporte para compressão de PDF
try:
    from PyPDF2 import PdfWriter, PdfReader
    import io
    PDF_COMPRESSION_AVAILABLE = True
except ImportError:
    PDF_COMPRESSION_AVAILABLE = False

class GeradorCurriculo:
    def __init__(self, dados_pessoais):
        self.dados = dados_pessoais
        self.doc = Document()
    
    def adicionar_cabecalho(self):
        self.doc.add_heading(self.dados['nome'], level=0)
        
        # Endereço
        endereco = f"Endereço: {self.dados['endereco']}"
        self.doc.add_paragraph(endereco)
        
        # Contatos
        contatos = f"Telefone: {self.dados['telefone']} | Email: {self.dados['email']} | LinkedIn: {self.dados['linkedin']}"
        self.doc.add_paragraph(contatos)
    
    def adicionar_objetivo(self):
        self.doc.add_heading("Objetivo", level=1)
        self.doc.add_paragraph(self.dados['objetivo'])
    
    def adicionar_experiencia(self):
        self.doc.add_heading("Experiência Profissional", level=1)
        for exp in self.dados['experiencias']:
            self.doc.add_heading(exp['empresa'], level=2)
            self.doc.add_paragraph(f"{exp['cargo']} | {exp['periodo']} | {exp['local']}")
            self.doc.add_paragraph(exp['descricao'])
    
    def adicionar_formacao(self):
        self.doc.add_heading("Formação Acadêmica", level=1)
        for formacao in self.dados['formacao']:
            self.doc.add_paragraph(formacao['curso'])
            self.doc.add_paragraph(f"{formacao['periodo']}\nStatus: {formacao['status']}\nTipo: {formacao['tipo']}\nCampus: {formacao['campus']}")
    
    def adicionar_skills(self):
        self.doc.add_heading("Skills", level=1)
        self.doc.add_paragraph(self.dados['skills'])
    
    def adicionar_competencias(self):
        self.doc.add_heading("Competências", level=1)
        self.doc.add_paragraph(self.dados['competencias'])
    
    def adicionar_cursos_certificacoes(self):
        self.doc.add_heading("Cursos e Certificações", level=1)
        for curso in self.dados['cursos_certificacoes']:
            self.doc.add_paragraph(f"• {curso}")
    
    def adicionar_informacoes_adicionais(self):
        self.doc.add_heading("Informações Adicionais", level=1)
        for info in self.dados['informacoes_adicionais']:
            self.doc.add_paragraph(f"• {info}")
    
    def adicionar_idiomas(self):
        self.doc.add_heading("Idiomas", level=1)
        for idioma in self.dados['idiomas']:
            self.doc.add_paragraph(f"• {idioma['idioma']}")
    
    def comprimir_pdf(self, caminho_pdf):
        """Comprime o PDF para reduzir o tamanho do arquivo"""
        if not PDF_COMPRESSION_AVAILABLE:
            return False
        
        try:
            # Ler o PDF original
            with open(caminho_pdf, 'rb') as file:
                reader = PdfReader(file)
                writer = PdfWriter()
                
                # Copiar todas as páginas
                for page in reader.pages:
                    # Comprimir a página
                    page.compress_content_streams()
                    writer.add_page(page)
                
                # Salvar o PDF comprimido
                with open(caminho_pdf, 'wb') as output_file:
                    writer.write(output_file)
            
            return True
        except Exception as e:
            print(f"Erro ao comprimir PDF: {e}")
            return False
    
    def gerar_curriculo(self, nome_arquivo_docx="Curriculo_Raimundo_Marques.docx", sobrescrever=True):
        try:
            pasta_destino = "Docs"
            if not os.path.exists(pasta_destino):
                os.makedirs(pasta_destino)

            caminho_docx = os.path.join(pasta_destino, nome_arquivo_docx)
            caminho_pdf = caminho_docx.replace(".docx", ".pdf")

            if not sobrescrever and (os.path.exists(caminho_docx) or os.path.exists(caminho_pdf)):
                resposta = input(f"Arquivos já existem em {pasta_destino}. Deseja sobrescrever? (s/n): ").lower()
                if resposta not in ['s', 'sim', 'y', 'yes']:
                    print("Operação cancelada.")
                    return False

            # Adiciona conteúdo no documento
            self.adicionar_cabecalho()
            self.adicionar_objetivo()
            self.adicionar_experiencia()
            self.adicionar_formacao()
            self.adicionar_cursos_certificacoes()
            self.adicionar_skills()
            self.adicionar_competencias()
            self.adicionar_informacoes_adicionais()
            self.adicionar_idiomas()

            # Salvar DOCX
            self.doc.save(caminho_docx)
            print(f"DOCX gerado com sucesso! Arquivo: {caminho_docx}")

            # Converter para PDF conforme o sistema operacional
            sistema = platform.system().lower()
            pdf_gerado = False
            
            if sistema == "windows":
                if docx2pdf_convert is not None:
                    try:
                        docx2pdf_convert(caminho_docx, caminho_pdf)
                        pdf_gerado = True
                        print(f"PDF gerado com sucesso! Arquivo: {caminho_pdf}")
                    except Exception as conv_err:
                        print(f"DOCX gerado, mas falha ao converter para PDF (docx2pdf): {conv_err}")
                else:
                    print("docx2pdf não está instalado. Instale com 'pip install docx2pdf'.")
            else:
                # Linux e outros sistemas
                try:
                    subprocess.run([
                        "libreoffice", "--headless", "--convert-to", "pdf", "--outdir", pasta_destino, caminho_docx
                    ], check=True)
                    pdf_gerado = True
                    print(f"PDF gerado com sucesso! Arquivo: {caminho_pdf}")
                except Exception as conv_err:
                    print(f"DOCX gerado, mas falha ao converter para PDF (libreoffice): {conv_err}")

            # Comprimir o PDF e garantir <= 4MB
            if pdf_gerado and os.path.exists(caminho_pdf):
                tamanho_antes = os.path.getsize(caminho_pdf) / (1024 * 1024)  # MB
                print(f"Tamanho inicial do PDF: {tamanho_antes:.2f}MB")

                try:
                    tmp_pdf = caminho_pdf + ".tmp.pdf"
                    subprocess.run([
                        "gs", "-sDEVICE=pdfwrite", "-dCompatibilityLevel=1.4",
                        "-dPDFSETTINGS=/ebook",   # /screen (mais leve), /ebook (qualidade boa)
                        "-dNOPAUSE", "-dQUIET", "-dBATCH",
                        f"-sOutputFile={tmp_pdf}", caminho_pdf
                    ], check=True)

                    tamanho_depois = os.path.getsize(tmp_pdf) / (1024 * 1024)
                    if tamanho_depois <= 4:
                        os.replace(tmp_pdf, caminho_pdf)  # substitui o original
                        reducao = ((tamanho_antes - tamanho_depois) / tamanho_antes) * 100
                        print(f"✅ PDF comprimido: {tamanho_antes:.2f}MB → {tamanho_depois:.2f}MB (redução {reducao:.1f}%)")
                    else:
                        print(f"⚠️ Mesmo após compressão, PDF tem {tamanho_depois:.2f}MB (> 4MB)")
                except Exception as e:
                    print(f"Erro ao tentar comprimir PDF com Ghostscript: {e}")

            return True
        except Exception as e:
            print(f"Erro ao gerar currículo: {e}")
            return False

   
def main():
    dados_pessoais = {
        'nome': "Raimundo Marques de Freitas Filho",
        'endereco': "Rua Henoch Reis nº 618, B da Paz - Manaus/AM - CEP: 69048-020",
        'telefone': "+55 97 98411-1260 (WhatsApp)",
        'email': "raimundo.marques.ff@gmail.com",
        'linkedin': "https://www.linkedin.com/in/raimundo-marques-filho-06478b108/",
        'objetivo': "Atuar como Full Stack Developer, aplicando minhas habilidades em desenvolvimento web e integração de sistemas para criar soluções eficientes e inovadoras. Busco contribuir em projetos desafiadores, utilizando tecnologias modernas tanto no front-end quando no backend e colaborar no desenvolvimento de aplicações que agreguem valor ao negócio.",
        'experiencias': [
            {
                'empresa': "Instituto Cal-Comp de Pesquisa e Inovação Tecnológica da Amazônia – ICCT",
                'cargo': "Desenvolvedor Full Stack II",
                'periodo': "Maio de 2024 – Atual",
                'local': "Manaus/AM",
                'descricao': "• Desenvolvimento e manutenção de aplicações full stack.\n• Implementação de novas funcionalidades e otimização de código.\n• Integração de sistemas e APIs para aprimorar fluxos de trabalho internos.\n• Colaboração com equipes multidisciplinares para desenvolvimento de soluções inovadoras.\n• Desenvolvimento de implementações web e web/mobile.\n• Criação de soluções para linhas de montagem de equipamentos de fábricas.\n• Tecnologias mais utilizadas: Docker, Docker-Compose, Ambiente Linux, Vue 3, NestJS, React, Bootstrap, MUI React, Tailwind, Javascript, Typescript, C# Blazor."
            },
            {
                'empresa': "Faculdade Metropolitana de Manaus - FAMETRO",
                'cargo': "Programador PHP PL",
                'periodo': "Setembro de 2021 - Dezembro de 2023",
                'local': "Manaus",
                'descricao': "• Responsável pela implementação de melhorias e manutenção de aplicações web em um portal de serviços legado, desenvolvido em PHP.\n• Integração de informações com a API do ERP TOTVS RM – Educacional, utilizando Laravel 9 e protocolo SOAP para comunicação (JSON/XML)."
            },
            {
                'empresa': "Estaleiro ERAM Embarcações",
                'cargo': "Estagiário em Sistemas ERP - Totvs RM",
                'periodo': "Janeiro de 2021 - Agosto de 2021",
                'local': "Manaus",
                'descricao': "• Suporte técnico e atendimento aos usuários dos módulos de Contabilidade, Estoque e RH do ERP TOTVS RM.\n• Desenvolvimento de relatórios com o RM Reports e consultas SQL para extração e análise de dados."
            },
            {
                'empresa': "Cartório do 3º Ofício de Registro de Imóveis de Manaus",
                'cargo': "Auxiliar de Cartório em Geral",
                'periodo': "Agosto de 2013 - Agosto de 2018",
                'local': "Manaus",
                'descricao': "• Cadastro de matrículas de imóveis, digitalização de documentos e processamento de averbação e ofícios.\n• Suporte técnico básico, incluindo manutenção de sistemas operacionais e gerenciamento de impressoras em rede via TCP/IP."
            }
        ],
        'formacao': [
            {
                'curso': "Pós-Graduado em Engenharia de Software – UNIFAVIP Wyden – Martha Falcão, Manaus",
                'periodo': "Janeiro de 2024 - Dezembro de 2025",
                'status': "Concluído",
                'tipo': "Especialização",
                'campus': "Polo Universitário – Caruaru – PE"
            },
            {
                'curso': "Graduação em Análise e Desenvolvimento de Sistemas – UNIFAVIP Wyden – Martha Falcão, Manaus",
                'periodo': "Janeiro de 2021 - Dezembro de 2023",
                'status': "Concluído",
                'tipo': "Tecnólogo",
                'campus': "Polo Universitário – Caruaru - PE"
            }
        ],
        'cursos_certificacoes': [
            "PHP do Básico ao Avançado - Udemy",
            "Javascript e Typescript do Básico ao Avançado - Udemy",
            "Certificação MCSA SQL Server Database Administration - Udemy",
            "Certificação Java - Udemy"
        ],
        'skills': (
            "• Linguagens & Frameworks: PHP (P.O.O, PHP Nativo, Laravel 9, AdonisJS), Java (Spring Boot), "
            "Javascript, TypeScript, Node.js (NestJS, integração via JDBC com ERP TOTVS RM – módulo educacional), "
            "Vue 3, React JS, C# / .NET (Blazor)\n"
            "• Banco de Dados: MySQL (Certificação 40h), PostgreSQL, SQL Server, MariaDB, Oracle 11.2 "
            "(execução em containers Docker)\n"
            "• Banco de Dados Relacionais: Experiência com modelagem, consultas otimizadas, integrações "
            "multiplataforma e manutenção de ambientes containerizados.\n"
            "• Ferramentas & DevOps: Docker, Docker-Compose, Git/GitHub (controle de versão, branching, "
            "pull requests, workflows e pipelines CI/CD em ambientes de homologação e produção)\n"
            "• APIs & Integrações: REST, SOAP, JSON (integrações entre aplicações corporativas e ERPs TOTVS RM)\n"
            "• Testes de Aplicação e API: Cypress, Insomnia, Postman\n"
            "• Sistemas Operacionais: Windows e Linux (principal)\n"
            "• Projetos de Sistemas: Experiência na concepção, desenvolvimento e manutenção de sistemas completos, "
            "com foco em escalabilidade, performance, segurança e boas práticas (Clean Code, SOLID e Design Patterns)."
        ),

        'competencias': (
            "• Ética, Resiliência, Organização e Planejamento\n"
            "• Pensamento analítico e resolução de problemas\n"
            "• Trabalho em equipe e comunicação interpessoal\n"
            "• Experiência com versionamento e colaboração via Git/GitHub (branches, merges, resolução de conflitos)\n"
            "• Vivência em pipelines CI/CD e integração contínua de aplicações em ambiente Dockerizado\n"
            "• Capacidade de adaptação a novas tecnologias e frameworks."
        ),
        'informacoes_adicionais': [
            "Documentação Completa e Atualizada",
            "CNH Cat B",
            "GitHub: https://github.com/RaimundoMarques",
            "LinkedIn: https://www.linkedin.com/in/raimundo-marques-filho-06478b108"
        ],
        'idiomas': [
            {
                'idioma': "Inglês técnico"
            }
        ]
    }
    
    gerador = GeradorCurriculo(dados_pessoais)
    gerador.gerar_curriculo("Curriculo_Raimundo_Marques.docx", sobrescrever=True)

if __name__ == "__main__":
    main()
