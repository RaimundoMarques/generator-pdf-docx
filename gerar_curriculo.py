from docx import Document
from docx2pdf import convert  # Biblioteca para converter DOCX em PDF
import os

class GeradorCurriculo:
    def __init__(self, dados_pessoais):
        self.dados = dados_pessoais
        self.doc = Document()
    
    def adicionar_cabecalho(self):
        self.doc.add_heading(self.dados['nome'], level=0)
        
        # Endereço
        endereco = f"📍 {self.dados['endereco']}"
        self.doc.add_paragraph(endereco)
        
        # Contatos
        contatos = f"📱 {self.dados['telefone']} | 📧 {self.dados['email']} | 🔗 {self.dados['linkedin']}"
        self.doc.add_paragraph(contatos)
    
    def adicionar_objetivo(self):
        self.doc.add_heading("Objetivo", level=1)
        self.doc.add_paragraph(self.dados['objetivo'])
    
    def adicionar_experiencia(self):
        self.doc.add_heading("Experiência Profissional", level=1)
        for exp in self.dados['experiencias']:
            self.doc.add_heading(exp['empresa'], level=2)
            self.doc.add_paragraph(f"{exp['cargo']} | {exp['periodo']} | {exp['local']}")
            self.doc.add_paragraph(exp['descricao'])
    
    def adicionar_formacao(self):
        self.doc.add_heading("Formação Acadêmica", level=1)
        for formacao in self.dados['formacao']:
            self.doc.add_paragraph(formacao['curso'])
            self.doc.add_paragraph(f"{formacao['periodo']}\nStatus: {formacao['status']}\nTipo: {formacao['tipo']}\nCampus: {formacao['campus']}")
    
    def adicionar_skills(self):
        self.doc.add_heading("Skills", level=1)
        self.doc.add_paragraph(self.dados['skills'])
    
    def adicionar_competencias(self):
        self.doc.add_heading("Competências", level=1)
        self.doc.add_paragraph(self.dados['competencias'])
    
    def adicionar_cursos_certificacoes(self):
        self.doc.add_heading("Cursos e Certificações", level=1)
        for curso in self.dados['cursos_certificacoes']:
            self.doc.add_paragraph(f"• {curso}")
    
    def adicionar_informacoes_adicionais(self):
        self.doc.add_heading("Informações Adicionais", level=1)
        for info in self.dados['informacoes_adicionais']:
            self.doc.add_paragraph(f"• {info}")
    
    def adicionar_idiomas(self):
        self.doc.add_heading("Idiomas", level=1)
        for idioma in self.dados['idiomas']:
            self.doc.add_paragraph(f"• {idioma['idioma']}: {idioma['nivel']}")
    
    def gerar_curriculo(self, nome_arquivo_docx="Curriculo.docx", sobrescrever=True):
        try:
            # Verificar se arquivos existem e perguntar se deve sobrescrever
            nome_arquivo_pdf = nome_arquivo_docx.replace(".docx", ".pdf")
            
            if not sobrescrever and (os.path.exists(nome_arquivo_docx) or os.path.exists(nome_arquivo_pdf)):
                resposta = input(f"Arquivos já existem. Deseja sobrescrever? (s/n): ").lower()
                if resposta not in ['s', 'sim', 'y', 'yes']:
                    print("Operação cancelada.")
                    return False
            
            self.adicionar_cabecalho()
            self.adicionar_objetivo()
            self.adicionar_experiencia()
            self.adicionar_formacao()
            self.adicionar_cursos_certificacoes()
            self.adicionar_skills()
            self.adicionar_competencias()
            self.adicionar_informacoes_adicionais()
            self.adicionar_idiomas()
            
            # Salvar DOCX
            self.doc.save(nome_arquivo_docx)
            print(f"DOCX gerado com sucesso! Arquivo: {nome_arquivo_docx}")
            
            # Converter para PDF
            convert(nome_arquivo_docx, nome_arquivo_pdf)
            print(f"PDF gerado com sucesso! Arquivo: {nome_arquivo_pdf}")
            
            return True
        except Exception as e:
            print(f"Erro ao gerar currículo: {e}")
            return False

def main():
    dados_pessoais = {
        'nome': "Raimundo Marques de Freitas Filho",
        'endereco': "Rua Henoch Reis nº 618, B da Paz - Manaus/AM - CEP: 69048-020",
        'telefone': "+55 97 98411-1260 (WhatsApp)",
        'email': "raimundo.marques.ff@gmail.com",
        'linkedin': "https://www.linkedin.com/in/raimundo-marques-filho-06478b108/",
        'objetivo': "Atuar como Full Stack Developer, aplicando minhas habilidades em desenvolvimento web e integração de sistemas para criar soluções eficientes e inovadoras. Busco contribuir em projetos desafiadores, utilizando tecnologias modernas tanto no front-end quando no backend e colaborar no desenvolvimento de aplicações que agreguem valor ao negócio.",
        'experiencias': [
            {
                'empresa': "Instituto Cal-Comp de Pesquisa e Inovação Tecnológica da Amazônia – ICCT",
                'cargo': "Desenvolvedor Full Stack II",
                'periodo': "Maio de 2024 – Atual",
                'local': "Manaus/AM",
                'descricao': "• Desenvolvimento e manutenção de aplicações full stack.\n• Implementação de novas funcionalidades e otimização de código.\n• Integração de sistemas e APIs para aprimorar fluxos de trabalho internos.\n• Colaboração com equipes multidisciplinares para desenvolvimento de soluções inovadoras.\n• Desenvolvimento de implementações web e web/mobile.\n• Criação de soluções para linhas de montagem de equipamentos de fábricas.\n• Tecnologias mais utilizadas: Docker, Docker-Compose, Ambiente Linux, Vue 3, NestJS, React, Bootstrap, MUI React, Tailwind, Javascript, Typescript, C# Blazor."
            },
            {
                'empresa': "Faculdade Metropolitana de Manaus - FAMETRO",
                'cargo': "Programador PHP PL",
                'periodo': "Setembro de 2021 - Dezembro de 2023",
                'local': "Manaus",
                'descricao': "• Responsável pela implementação de melhorias e manutenção de aplicações web em um portal de serviços legado, desenvolvido em PHP.\n• Integração de informações com a API do ERP TOTVS RM – Educacional, utilizando Laravel 9 e protocolo SOAP para comunicação (JSON/XML)."
            },
            {
                'empresa': "Estaleiro ERAM Embarcações",
                'cargo': "Estagiário em Sistemas ERP - Totvs RM",
                'periodo': "Janeiro de 2021 - Agosto de 2021",
                'local': "Manaus",
                'descricao': "• Suporte técnico e atendimento aos usuários dos módulos de Contabilidade, Estoque e RH do ERP TOTVS RM.\n• Desenvolvimento de relatórios com o RM Reports e consultas SQL para extração e análise de dados."
            },
            {
                'empresa': "Cartório do 3º Ofício de Registro de Imóveis de Manaus",
                'cargo': "Auxiliar de Cartório em Geral",
                'periodo': "Agosto de 2013 - Agosto de 2018",
                'local': "Manaus",
                'descricao': "• Cadastro de matrículas de imóveis, digitalização de documentos e processamento de averbação e ofícios.\n• Suporte técnico básico, incluindo manutenção de sistemas operacionais e gerenciamento de impressoras em rede via TCP/IP."
            }
        ],
        'formacao': [
            {
                'curso': "Pós-Graduando em Engenharia de Software – UNIFAVIP Wyden – Martha Falcão, Manaus",
                'periodo': "Janeiro de 2024 - Dezembro de 2025 (Cursando – atualmente no 3º módulo de 4)",
                'status': "Cursando",
                'tipo': "Especialização",
                'campus': "Polo Universitário – Caruaru – PE"
            },
            {
                'curso': "Graduação em Análise e Desenvolvimento de Sistemas – UNIFAVIP Wyden – Martha Falcão, Manaus",
                'periodo': "Janeiro de 2021 - Dezembro de 2023",
                'status': "Concluído",
                'tipo': "Tecnólogo",
                'campus': "Polo Universitário – Caruaru - PE"
            }
        ],
        'cursos_certificacoes': [
            "PHP do Básico ao Avançado - Udemy",
            "Javascript e Typescript do Básico ao Avançado - Udemy",
            "Certificação MCSA SQL Server Database Administration - Udemy",
            "Certificação Java - Udemy"
        ],
        'skills': "• Linguagens & Frameworks: PHP (P.O.O, PHP Nativo, Laravel 9, AdonisJS), Javascript, TypeScript, Node.js, NestJS, Vue 3, React JS, C# Blazor\n• Banco de Dados: MySQL (Certificação 40h), PostgreSQL, SQL Server, MariaDB, Oracle\n• Ferramentas & DevOps: Docker, Docker-Compose, Git/GitHub (controle de versão, branching, pull requests e workflows com CI/CD)\n• APIs & Integrações: REST, SOAP, JSON\n• Testes de API: Insomnia, Postman\n• Sistemas Operacionais: Windows e Linux (principal)\n• Projetos de Sistemas: Experiência na concepção, desenvolvimento e manutenção de sistemas completos, garantindo escalabilidade, performance e boas práticas.",
        'competencias': "• Ética, Resiliência, Organização e Planejamento\n• Pensamento analítico e resolução de problemas\n• Trabalho em equipe\n• Experiência com controle de versão utilizando Git e GitHub, incluindo boas práticas de versionamento, resolução de conflitos, uso de branches e colaboração em equipe.",
        'informacoes_adicionais': [
            "Documentação Completa e Atualizada",
            "CNH Cat B",
            "GitHub: https://github.com/RaimundoMarques",
            "LinkedIn: https://www.linkedin.com/in/raimundo-marques-filho-06478b108"
        ],
        'idiomas': [
            {
                'idioma': "Inglês",
                'nivel': "Básico / Intermediário (Cursando)"
            }
        ]
    }
    
    gerador = GeradorCurriculo(dados_pessoais)
    gerador.gerar_curriculo("Curriculo_Raimundo_Marques.docx", sobrescrever=True)

if __name__ == "__main__":
    main()
